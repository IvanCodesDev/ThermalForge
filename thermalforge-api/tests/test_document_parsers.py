from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import fitz
import pytest
from docx import Document
from PIL import Image

from app.documents.ocr import OcrResult
from app.documents.parsers import DocumentParserRegistry
from app.documents.validation import DocumentValidator
from app.domain.errors import EncryptedDocument, InvalidDocument


class FakeOcrProvider:
    async def extract_text(self, image: bytes) -> OcrResult:
        assert image
        return OcrResult(text="OCR 热源标签 85°C", confidence=0.98)


def validator() -> DocumentValidator:
    return DocumentValidator(
        max_upload_bytes=20 * 1024 * 1024,
        max_archive_entries=2_000,
        max_archive_uncompressed_bytes=100 * 1024 * 1024,
        max_image_pixels=50_000_000,
    )


@pytest.mark.asyncio
async def test_parses_pdf_pages_with_source_references(tmp_path: Path) -> None:
    pdf_path = tmp_path / "joint.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Joint power 120 W\nAmbient temperature 25 C")
    pdf.save(pdf_path)
    pdf.close()

    document = validator().validate(
        path=pdf_path,
        original_filename="joint.pdf",
        declared_mime="application/pdf",
    )
    parsed = await DocumentParserRegistry(FakeOcrProvider()).parse(
        document,
        source_artifact_id="artifact-pdf",
    )

    assert parsed.page_count == 1
    assert "Joint power 120 W" in parsed.chunks[0].text
    assert parsed.chunks[0].page_number == 1
    assert parsed.chunks[0].source_artifact_id == "artifact-pdf"


@pytest.mark.asyncio
async def test_parses_docx_headings_and_tables(tmp_path: Path) -> None:
    docx_path = tmp_path / "constraints.docx"
    docx = Document()
    docx.add_heading("安装约束", level=1)
    docx.add_paragraph("外壳必须保持可拆卸。")
    table = docx.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "功率"
    table.cell(0, 1).text = "120 W"
    table.cell(1, 0).text = "环境温度"
    table.cell(1, 1).text = "25°C"
    docx.save(docx_path)

    document = validator().validate(
        path=docx_path,
        original_filename="constraints.docx",
        declared_mime=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    parsed = await DocumentParserRegistry(FakeOcrProvider()).parse(
        document,
        source_artifact_id="artifact-docx",
    )

    assert any(chunk.section_path == ["安装约束"] for chunk in parsed.chunks)
    assert parsed.tables[0].rows[0] == ["功率", "120 W"]


@pytest.mark.asyncio
async def test_extracts_and_ocr_scans_docx_images(tmp_path: Path) -> None:
    docx_path = tmp_path / "drawing.docx"
    image_buffer = BytesIO()
    Image.new("RGB", (240, 160), color="white").save(image_buffer, format="PNG")
    image_buffer.seek(0)
    docx = Document()
    docx.add_picture(image_buffer)
    docx.save(docx_path)

    document = validator().validate(
        path=docx_path,
        original_filename="drawing.docx",
        declared_mime=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    parsed = await DocumentParserRegistry(FakeOcrProvider()).parse(
        document,
        source_artifact_id="artifact-docx-image",
    )

    assert parsed.images[0].width == 240
    assert parsed.images[0].height == 160
    assert parsed.chunks[0].content_type == "ocr"


@pytest.mark.asyncio
async def test_uses_ocr_for_images_and_preserves_dimensions(tmp_path: Path) -> None:
    image_path = tmp_path / "drawing.png"
    image = Image.new("RGB", (640, 480), color="white")
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    image_path.write_bytes(buffer.getvalue())

    document = validator().validate(
        path=image_path,
        original_filename="drawing.png",
        declared_mime="image/png",
    )
    parsed = await DocumentParserRegistry(FakeOcrProvider()).parse(
        document,
        source_artifact_id="artifact-image",
    )

    assert parsed.images[0].width == 640
    assert parsed.images[0].height == 480
    assert parsed.chunks[0].text == "OCR 热源标签 85°C"


@pytest.mark.asyncio
async def test_uses_ocr_for_scanned_pdf_pages(tmp_path: Path) -> None:
    image_buffer = BytesIO()
    Image.new("RGB", (300, 120), color="white").save(image_buffer, format="PNG")
    pdf_path = tmp_path / "scan.pdf"
    pdf = fitz.open()
    page = pdf.new_page(width=300, height=120)
    page.insert_image(page.rect, stream=image_buffer.getvalue())
    pdf.save(pdf_path)
    pdf.close()

    document = validator().validate(
        path=pdf_path,
        original_filename="scan.pdf",
        declared_mime="application/pdf",
    )
    parsed = await DocumentParserRegistry(FakeOcrProvider()).parse(
        document,
        source_artifact_id="artifact-scan",
    )

    assert parsed.chunks[0].content_type == "ocr"
    assert parsed.chunks[0].page_number == 1
    assert parsed.images[0].ocr_confidence == 0.98


@pytest.mark.asyncio
async def test_extracts_images_from_text_pdf_pages(tmp_path: Path) -> None:
    image_buffer = BytesIO()
    Image.new("RGB", (120, 80), color="white").save(image_buffer, format="PNG")
    pdf_path = tmp_path / "mixed.pdf"
    pdf = fitz.open()
    page = pdf.new_page(width=400, height=300)
    page.insert_text((20, 30), "Assembly drawing")
    page.insert_image(fitz.Rect(20, 60, 140, 140), stream=image_buffer.getvalue())
    pdf.save(pdf_path)
    pdf.close()

    document = validator().validate(
        path=pdf_path,
        original_filename="mixed.pdf",
        declared_mime="application/pdf",
    )
    parsed = await DocumentParserRegistry(FakeOcrProvider()).parse(
        document,
        source_artifact_id="artifact-mixed",
    )

    assert parsed.images[0].width == 120
    assert any(chunk.content_type == "ocr" for chunk in parsed.chunks)


@pytest.mark.parametrize(
    ("extension", "mime_type", "image_format"),
    [
        (".jpg", "image/jpeg", "JPEG"),
        (".webp", "image/webp", "WEBP"),
    ],
)
@pytest.mark.asyncio
async def test_accepts_supported_image_variants(
    tmp_path: Path,
    extension: str,
    mime_type: str,
    image_format: str,
) -> None:
    image_path = tmp_path / f"drawing{extension}"
    Image.new("RGB", (80, 60), color="white").save(
        image_path,
        format=image_format,
    )

    document = validator().validate(
        path=image_path,
        original_filename=image_path.name,
        declared_mime=mime_type,
    )
    parsed = await DocumentParserRegistry(FakeOcrProvider()).parse(
        document,
        source_artifact_id=f"artifact-{image_format.lower()}",
    )

    assert parsed.images[0].width == 80
    assert parsed.chunks[0].content_type == "ocr"


@pytest.mark.asyncio
async def test_accepts_plain_text_with_detected_encoding(tmp_path: Path) -> None:
    text_path = tmp_path / "constraints.txt"
    text_path.write_text("环境温度 25°C\n最大功率 120 W", encoding="utf-8")

    document = validator().validate(
        path=text_path,
        original_filename=text_path.name,
        declared_mime="text/plain",
    )
    parsed = await DocumentParserRegistry(FakeOcrProvider()).parse(
        document,
        source_artifact_id="artifact-text",
    )

    assert "最大功率 120 W" in parsed.chunks[0].text


def test_rejects_encrypted_pdf_and_corrupt_docx(tmp_path: Path) -> None:
    encrypted_path = tmp_path / "encrypted.pdf"
    pdf = fitz.open()
    pdf.new_page()
    encrypted_path.write_bytes(
        pdf.tobytes(
            encryption=fitz.PDF_ENCRYPT_AES_256,
            owner_pw="owner",
            user_pw="secret",
        )
    )
    pdf.close()
    corrupt_docx = tmp_path / "corrupt.docx"
    corrupt_docx.write_bytes(b"PK\x03\x04not-a-real-docx")

    with pytest.raises(EncryptedDocument):
        validator().validate(
            path=encrypted_path,
            original_filename="encrypted.pdf",
            declared_mime="application/pdf",
        )

    with pytest.raises(InvalidDocument):
        validator().validate(
            path=corrupt_docx,
            original_filename="corrupt.docx",
            declared_mime=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )


def test_rejects_empty_macro_enabled_and_oversized_image_files(
    tmp_path: Path,
) -> None:
    empty_text = tmp_path / "empty.txt"
    empty_text.touch()
    macro_docx = tmp_path / "macro.docx"
    with ZipFile(macro_docx, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types />")
        archive.writestr("word/document.xml", "<document />")
        archive.writestr("word/vbaProject.bin", b"macro")
    large_image = tmp_path / "large.png"
    Image.new("RGB", (100, 100), color="white").save(large_image)

    with pytest.raises(InvalidDocument, match="empty"):
        validator().validate(
            path=empty_text,
            original_filename="empty.txt",
            declared_mime="text/plain",
        )
    with pytest.raises(InvalidDocument, match="Macro-enabled"):
        validator().validate(
            path=macro_docx,
            original_filename="macro.docx",
            declared_mime=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )
    with pytest.raises(InvalidDocument, match="dimensions"):
        DocumentValidator(
            max_upload_bytes=20 * 1024 * 1024,
            max_archive_entries=2_000,
            max_archive_uncompressed_bytes=100 * 1024 * 1024,
            max_image_pixels=9_999,
        ).validate(
            path=large_image,
            original_filename="large.png",
            declared_mime="image/png",
        )
