import asyncio
import unicodedata
from io import BytesIO
from pathlib import Path

import fitz
from charset_normalizer import from_bytes
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image

from app.documents.chunking import TextSegment, build_chunks
from app.documents.ocr import OcrProvider
from app.documents.schemas import (
    DocumentChunk,
    DocumentImage,
    DocumentTable,
    ParsedDocument,
)
from app.documents.validation import ValidatedDocument
from app.domain.errors import InvalidDocument


class DocumentParserRegistry:
    def __init__(
        self,
        ocr_provider: OcrProvider,
        *,
        max_chunk_chars: int = 1_600,
        overlap_chars: int = 200,
    ) -> None:
        self._ocr_provider = ocr_provider
        self._max_chunk_chars = max_chunk_chars
        self._overlap_chars = overlap_chars

    def _chunks(
        self,
        source_artifact_id: str,
        segments: list[TextSegment],
    ) -> list[DocumentChunk]:
        return build_chunks(
            source_artifact_id=source_artifact_id,
            segments=segments,
            max_chars=self._max_chunk_chars,
            overlap_chars=self._overlap_chars,
        )

    @staticmethod
    def _decode_text(path: Path) -> str:
        match = from_bytes(path.read_bytes()).best()
        if match is None:
            raise InvalidDocument("Text encoding could not be detected.")
        normalized = unicodedata.normalize("NFC", str(match))
        return normalized.replace("\r\n", "\n").replace("\r", "\n")

    @staticmethod
    def _markdown_segments(text: str) -> list[TextSegment]:
        sections: list[str] = []
        segments: list[TextSegment] = []
        paragraph: list[str] = []

        def flush_paragraph() -> None:
            content = "\n".join(paragraph).strip()
            if content:
                segments.append(
                    TextSegment(
                        text=content,
                        page_number=1,
                        section_path=tuple(sections),
                    )
                )
            paragraph.clear()

        for line in text.splitlines():
            stripped = line.strip()
            heading_level = len(stripped) - len(stripped.lstrip("#"))
            if 1 <= heading_level <= 6 and stripped[heading_level:].startswith(" "):
                flush_paragraph()
                heading = stripped[heading_level:].strip()
                sections[:] = sections[: heading_level - 1]
                sections.append(heading)
                continue
            if not stripped:
                flush_paragraph()
            else:
                paragraph.append(stripped)
        flush_paragraph()
        return segments

    async def _parse_text(
        self,
        document: ValidatedDocument,
        source_artifact_id: str,
    ) -> ParsedDocument:
        text = await asyncio.to_thread(self._decode_text, document.path)
        segments = (
            self._markdown_segments(text)
            if document.extension == ".md"
            else [TextSegment(text=text, page_number=1)]
        )
        chunks = self._chunks(source_artifact_id, segments)
        return ParsedDocument(
            page_count=1,
            chunks=chunks,
            warnings=[] if chunks else ["No readable text was found."],
        )

    async def _parse_pdf(
        self,
        document: ValidatedDocument,
        source_artifact_id: str,
    ) -> ParsedDocument:
        pdf = fitz.open(document.path)
        page_count = pdf.page_count
        segments: list[TextSegment] = []
        tables: list[DocumentTable] = []
        images: list[DocumentImage] = []
        warnings: list[str] = []

        try:
            for page_index, page in enumerate(pdf):
                page_number = page_index + 1
                embedded_ocr_found = False
                for image_index, image_info in enumerate(
                    page.get_images(full=True)
                ):
                    extracted = pdf.extract_image(int(image_info[0]))
                    image_payload = extracted["image"]
                    with Image.open(BytesIO(image_payload)) as embedded_image:
                        width, height = embedded_image.size
                    ocr = await self._ocr_provider.extract_text(image_payload)
                    images.append(
                        DocumentImage(
                            id=(
                                f"{source_artifact_id}:page:{page_number}:"
                                f"image:{image_index}"
                            ),
                            source_artifact_id=source_artifact_id,
                            width=width,
                            height=height,
                            page_number=page_number,
                            ocr_confidence=ocr.confidence,
                        )
                    )
                    if ocr.text:
                        embedded_ocr_found = True
                        segments.append(
                            TextSegment(
                                text=ocr.text,
                                page_number=page_number,
                                content_type="ocr",
                            )
                        )

                text = page.get_text("text").strip()
                if text:
                    segments.append(TextSegment(text=text, page_number=page_number))
                elif not embedded_ocr_found:
                    pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    rendered = pixmap.tobytes("png")
                    ocr = await self._ocr_provider.extract_text(rendered)
                    images.append(
                        DocumentImage(
                            id=f"{source_artifact_id}:page-image:{page_number}",
                            source_artifact_id=source_artifact_id,
                            width=pixmap.width,
                            height=pixmap.height,
                            page_number=page_number,
                            ocr_confidence=ocr.confidence,
                        )
                    )
                    if ocr.text:
                        segments.append(
                            TextSegment(
                                text=ocr.text,
                                page_number=page_number,
                                content_type="ocr",
                            )
                        )
                    else:
                        warnings.append(f"Page {page_number} has no readable text.")

                try:
                    located_tables = page.find_tables()
                except (AttributeError, RuntimeError):
                    located_tables = None
                if located_tables is not None:
                    for table_index, table in enumerate(located_tables.tables):
                        rows = [
                            [str(cell or "").strip() for cell in row]
                            for row in table.extract()
                        ]
                        if rows:
                            tables.append(
                                DocumentTable(
                                    id=(
                                        f"{source_artifact_id}:page:{page_number}:"
                                        f"table:{table_index}"
                                    ),
                                    source_artifact_id=source_artifact_id,
                                    page_number=page_number,
                                    rows=rows,
                                )
                            )
        finally:
            pdf.close()

        return ParsedDocument(
            page_count=page_count,
            chunks=self._chunks(source_artifact_id, segments),
            tables=tables,
            images=images,
            warnings=warnings,
        )

    @staticmethod
    def _read_docx(
        path: Path,
    ) -> tuple[list[TextSegment], list[list[list[str]]], list[bytes]]:
        docx = Document(str(path))
        sections: list[str] = []
        segments: list[TextSegment] = []
        tables: list[list[list[str]]] = []

        for block in docx.iter_inner_content():
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if not text:
                    continue
                style_name = block.style.name if block.style is not None else ""
                if style_name.startswith("Heading "):
                    try:
                        level = int(style_name.removeprefix("Heading "))
                    except ValueError:
                        level = 1
                    sections[:] = sections[: level - 1]
                    sections.append(text)
                else:
                    segments.append(
                        TextSegment(
                            text=text,
                            page_number=1,
                            section_path=tuple(sections),
                        )
                    )
            elif isinstance(block, Table):
                rows = [
                    [cell.text.strip() for cell in row.cells]
                    for row in block.rows
                ]
                if rows:
                    tables.append(rows)
        image_payloads = [
            relationship.target_part.blob
            for relationship in docx.part.rels.values()
            if relationship.reltype == RELATIONSHIP_TYPE.IMAGE
        ]
        return segments, tables, image_payloads

    async def _parse_docx(
        self,
        document: ValidatedDocument,
        source_artifact_id: str,
    ) -> ParsedDocument:
        try:
            segments, raw_tables, image_payloads = await asyncio.to_thread(
                self._read_docx,
                document.path,
            )
        except (KeyError, ValueError, OSError) as error:
            raise InvalidDocument("DOCX content could not be parsed.") from error

        tables = [
            DocumentTable(
                id=f"{source_artifact_id}:table:{index}",
                source_artifact_id=source_artifact_id,
                page_number=1,
                rows=rows,
            )
            for index, rows in enumerate(raw_tables)
        ]
        images: list[DocumentImage] = []
        for index, payload in enumerate(image_payloads):
            with Image.open(BytesIO(payload)) as image:
                width, height = image.size
            ocr = await self._ocr_provider.extract_text(payload)
            images.append(
                DocumentImage(
                    id=f"{source_artifact_id}:image:{index}",
                    source_artifact_id=source_artifact_id,
                    width=width,
                    height=height,
                    page_number=1,
                    ocr_confidence=ocr.confidence,
                )
            )
            if ocr.text:
                segments.append(
                    TextSegment(
                        text=ocr.text,
                        page_number=1,
                        content_type="ocr",
                    )
                )
        chunks = self._chunks(source_artifact_id, segments)
        return ParsedDocument(
            page_count=1,
            chunks=chunks,
            tables=tables,
            images=images,
            warnings=(
                []
                if chunks or tables or images
                else ["DOCX contains no readable content."]
            ),
        )

    async def _parse_image(
        self,
        document: ValidatedDocument,
        source_artifact_id: str,
    ) -> ParsedDocument:
        payload = await asyncio.to_thread(document.path.read_bytes)
        with Image.open(BytesIO(payload)) as image:
            width, height = image.size
        ocr = await self._ocr_provider.extract_text(payload)
        segments = (
            [
                TextSegment(
                    text=ocr.text,
                    page_number=1,
                    content_type="ocr",
                )
            ]
            if ocr.text
            else []
        )
        return ParsedDocument(
            page_count=1,
            chunks=self._chunks(source_artifact_id, segments),
            images=[
                DocumentImage(
                    id=f"{source_artifact_id}:image:0",
                    source_artifact_id=source_artifact_id,
                    width=width,
                    height=height,
                    page_number=1,
                    ocr_confidence=ocr.confidence,
                )
            ],
            warnings=[] if ocr.text else ["OCR did not detect text in the image."],
        )

    async def parse(
        self,
        document: ValidatedDocument,
        *,
        source_artifact_id: str,
    ) -> ParsedDocument:
        if document.extension in {".txt", ".md"}:
            return await self._parse_text(document, source_artifact_id)
        if document.extension == ".pdf":
            return await self._parse_pdf(document, source_artifact_id)
        if document.extension == ".docx":
            return await self._parse_docx(document, source_artifact_id)
        return await self._parse_image(document, source_artifact_id)
