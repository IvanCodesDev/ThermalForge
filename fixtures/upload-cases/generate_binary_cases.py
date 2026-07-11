from pathlib import Path
import textwrap

import fitz
from docx import Document


ROOT = Path(__file__).resolve().parent


def generate_pdf() -> None:
    source = (ROOT / "01-foc-joint-complete.md").read_text(encoding="utf-8")
    lines: list[str] = []
    for raw_line in source.splitlines():
        line = raw_line.lstrip("# ").strip()
        if not line:
            lines.append("")
            continue
        lines.extend(textwrap.wrap(line, width=42, break_long_words=False))

    document = fitz.open()
    for offset in range(0, len(lines), 34):
        page = document.new_page(width=595, height=842)
        y = 54
        for line in lines[offset : offset + 34]:
            page.insert_text(
                (48, y),
                line,
                fontname="china-s",
                fontsize=11,
                color=(0.08, 0.08, 0.08),
            )
            y += 21
    document.save(ROOT / "04-foc-joint-complete.pdf", garbage=4, deflate=True)
    document.close()


def generate_docx() -> None:
    source = (ROOT / "02-liquid-cold-plate-complete.txt").read_text(
        encoding="utf-8"
    )
    document = Document()
    for index, line in enumerate(source.splitlines()):
        value = line.strip()
        if not value:
            document.add_paragraph()
        elif index == 0:
            document.add_heading(value, level=1)
        elif value.endswith("："):
            document.add_heading(value.rstrip("："), level=2)
        else:
            document.add_paragraph(value)
    document.save(ROOT / "05-liquid-cold-plate-complete.docx")


if __name__ == "__main__":
    generate_pdf()
    generate_docx()
